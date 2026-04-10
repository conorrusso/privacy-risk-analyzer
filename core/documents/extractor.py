"""
Bandit Document Extractor
==========================
Raw text extraction from supported file formats.
Never crashes on a single file failure — logs and continues.

Supported: PDF, DOCX, DOC, HTML, HTM, TXT, MD, JSON
Skipped silently: XLSX, CSV, ZIP, PNG, JPG, JPEG, GIF, MP4, MOV
"""
from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path

# pdfminer emits harmless FontBBox warnings for malformed fonts in some PDFs
logging.getLogger("pdfminer").setLevel(logging.ERROR)

SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".doc",
    ".html", ".htm", ".txt", ".md", ".json"
}

SKIP_EXTENSIONS = {
    ".xlsx", ".csv", ".zip", ".png", ".jpg",
    ".jpeg", ".gif", ".mp4", ".mov"
}


@dataclass
class ExtractionResult:
    file_path: str
    file_name: str
    format: str         # pdf | docx | html | txt | json
    text: str
    char_count: int
    page_count: int     # PDFs only, else 0
    extraction_ok: bool
    error: str | None


class DocumentExtractor:

    def extract(self, file_path: str) -> ExtractionResult:
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext in SKIP_EXTENSIONS:
            return ExtractionResult(
                file_path=str(path),
                file_name=path.name,
                format=ext.lstrip("."),
                text="",
                char_count=0,
                page_count=0,
                extraction_ok=False,
                error=f"Unsupported format {ext}"
            )

        try:
            if ext == ".pdf":
                return self._extract_pdf(path)
            elif ext in (".docx", ".doc"):
                return self._extract_docx(path)
            elif ext in (".html", ".htm"):
                return self._extract_html(path)
            elif ext in (".txt", ".md"):
                return self._extract_text(path)
            elif ext == ".json":
                return self._extract_json(path)
            else:
                return ExtractionResult(
                    file_path=str(path),
                    file_name=path.name,
                    format=ext.lstrip("."),
                    text="",
                    char_count=0,
                    page_count=0,
                    extraction_ok=False,
                    error=(
                        f"Unsupported format {ext}. "
                        f"Supported: PDF DOCX HTML TXT MD JSON"
                    )
                )
        except Exception as e:
            return ExtractionResult(
                file_path=str(path),
                file_name=path.name,
                format=ext.lstrip("."),
                text="",
                char_count=0,
                page_count=0,
                extraction_ok=False,
                error=str(e)
            )

    def _extract_pdf(self, path: Path) -> ExtractionResult:
        import pdfplumber
        pages = []
        page_count = 0

        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)

        full_text = "\n\n".join(pages).strip()

        # Detect scanned PDF (no text layer)
        if len(full_text) < 200 and page_count > 0:
            return ExtractionResult(
                file_path=str(path),
                file_name=path.name,
                format="pdf",
                text="",
                char_count=0,
                page_count=page_count,
                extraction_ok=False,
                error=(
                    f"Scanned PDF detected — text layer absent. "
                    f"OCR support coming in v1.2. "
                    f"Try a text-based PDF version."
                )
            )

        return ExtractionResult(
            file_path=str(path),
            file_name=path.name,
            format="pdf",
            text=full_text,
            char_count=len(full_text),
            page_count=page_count,
            extraction_ok=True,
            error=None
        )

    def _extract_docx(self, path: Path) -> ExtractionResult:
        from docx import Document
        doc = Document(str(path))
        parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        # Extract table cells — important for DPAs
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())

        full_text = "\n\n".join(parts).strip()
        return ExtractionResult(
            file_path=str(path),
            file_name=path.name,
            format="docx",
            text=full_text,
            char_count=len(full_text),
            page_count=0,
            extraction_ok=True,
            error=None
        )

    def _extract_html(self, path: Path) -> ExtractionResult:
        from core.tools.parse import html_to_text
        raw = path.read_text(encoding="utf-8", errors="replace")
        text = html_to_text(raw)
        return ExtractionResult(
            file_path=str(path),
            file_name=path.name,
            format="html",
            text=text,
            char_count=len(text),
            page_count=0,
            extraction_ok=True,
            error=None
        )

    def _extract_text(self, path: Path) -> ExtractionResult:
        text = path.read_text(encoding="utf-8", errors="replace").strip()
        return ExtractionResult(
            file_path=str(path),
            file_name=path.name,
            format=path.suffix.lstrip("."),
            text=text,
            char_count=len(text),
            page_count=0,
            extraction_ok=True,
            error=None
        )

    def _extract_json(self, path: Path) -> ExtractionResult:
        import json
        raw = json.loads(
            path.read_text(encoding="utf-8", errors="replace")
        )

        def flatten(obj, parts=None):
            if parts is None:
                parts = []
            if isinstance(obj, str):
                parts.append(obj)
            elif isinstance(obj, dict):
                for v in obj.values():
                    flatten(v, parts)
            elif isinstance(obj, list):
                for item in obj:
                    flatten(item, parts)
            return parts

        text = "\n".join(flatten(raw)).strip()
        return ExtractionResult(
            file_path=str(path),
            file_name=path.name,
            format="json",
            text=text,
            char_count=len(text),
            page_count=0,
            extraction_ok=True,
            error=None
        )
